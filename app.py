from __future__ import annotations

import hashlib
import json
import logging
import os
import re
import secrets
import sqlite3
import time
import uuid
from datetime import datetime, timedelta, timezone
from typing import Any, Dict, List, Tuple

from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)

if not os.getenv("OPENAI_API_KEY"):
    logger.warning(
        "OPENAI_API_KEY is not set. Extraction requests will fail until it is "
        "added to .env or the environment."
    )


from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    session,
    flash,
    make_response,
)
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

try:
    from pypdf import PdfReader  # type: ignore
except ModuleNotFoundError:  # pragma: no cover - handled at runtime
    PdfReader = None

try:
    import docx  # type: ignore - python-docx
except ModuleNotFoundError:  # pragma: no cover - handled at runtime
    docx = None

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY") or secrets.token_hex(32)
app.config["MAX_CONTENT_LENGTH"] = int(os.environ.get("MAX_UPLOAD_MB", "10")) * 1024 * 1024
app.config["EXTRACT_RATE_LIMIT"] = os.environ.get("EXTRACT_RATE_LIMIT", "20 per hour")

# Trust X-Forwarded-* headers when running behind a reverse proxy (Render, Railway, ...)
# so rate limiting sees the real client IP instead of the proxy's.
if os.environ.get("TRUST_PROXY") or os.environ.get("RENDER") or os.environ.get("RAILWAY_ENVIRONMENT"):
    from werkzeug.middleware.proxy_fix import ProxyFix

    app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1)

limiter = Limiter(
    get_remote_address,
    app=app,
    storage_uri=os.environ.get("RATELIMIT_STORAGE_URI", "memory://"),
)

# --- Persistence (SQLite) ---
DB_PATH = os.environ.get("SYLLABUS_DB_PATH") or os.path.join(os.path.dirname(__file__), "syllabus.db")

def _db_conn():
    return sqlite3.connect(DB_PATH)

def _init_db():
    with _db_conn() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS saved_entry (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                visitor_id TEXT NOT NULL,
                created_at TEXT NOT NULL,
                term_year INTEGER NULL,
                start_month INTEGER NULL,
                end_month INTEGER NULL,
                events_json TEXT NOT NULL,
                recurring_json TEXT NOT NULL,
                label TEXT NULL
            );
            """
        )
        conn.execute("CREATE INDEX IF NOT EXISTS idx_saved_entry_visitor_created ON saved_entry(visitor_id, created_at DESC);")
        # Cache table for extraction results
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS extraction_cache (
                cache_key TEXT PRIMARY KEY,
                created_at TEXT NOT NULL,
                params_json TEXT NOT NULL,
                text_len INTEGER NOT NULL,
                events_json TEXT NOT NULL
            );
            """
        )

_init_db()


def _utcnow_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")


def _get_or_create_visitor_id() -> str:
    vid = request.cookies.get("visitor_id")
    if vid and isinstance(vid, str) and len(vid) <= 64:
        return vid
    return uuid.uuid4().hex

@app.after_request
def _ensure_visitor_cookie(resp):
    if not request.cookies.get("visitor_id"):
        resp.set_cookie("visitor_id", _get_or_create_visitor_id(), httponly=True, samesite="Lax", max_age=60*60*24*365*2)
    return resp


@app.route("/", methods=["GET", "POST"])
@limiter.limit(lambda: app.config["EXTRACT_RATE_LIMIT"], methods=["POST"])
def index():
    error = None
    events: List[Dict[str, Any]] = []
    recurring_items: List[Dict[str, Any]] = []

    if request.method == "POST":
        uploaded = request.files.get("syllabus")
        if not uploaded or uploaded.filename == "":
            error = "Please choose a syllabus file to upload."
        else:
            try:
                year_str = (request.form.get("year") or "").strip()
                year = int(year_str) if year_str.isdigit() else None
                sm_str = (request.form.get("start_month") or "").strip()
                em_str = (request.form.get("end_month") or "").strip()
                start_month = int(sm_str) if sm_str.isdigit() else None
                end_month = int(em_str) if em_str.isdigit() else None

                text = extract_text(uploaded)
                due_dates = extract_due_dates(
                    text,
                    user_year=year,
                    start_month=start_month,
                    end_month=end_month,
                )
                # Remember user parameters for Save Entry
                session["term_year"] = year
                session["start_month"] = start_month
                session["end_month"] = end_month
                # Split dated and recurring items
                dated_events, recurring_items = split_events(due_dates)
                # Store only dated events for ICS usage
                session["events"] = [
                    {"date": event["date"].isoformat(), "description": event["description"]}
                    for event in dated_events
                ]
                # Store recurring items separately for display-only
                session["recurring"] = [
                    {"date": None, "description": item["description"]}
                    for item in recurring_items
                ]
                # Only build calendar links for dated events
                events = add_calendar_links(dated_events)
                if not events:
                    flash("No due dates were detected. Try cleaning up the PDF or uploading a text export.", "info")
            except RuntimeError as exc:
                error = str(exc)
            except Exception:  # pragma: no cover - generic failure handler
                logger.exception("Failed to process uploaded syllabus")
                error = "We couldn't read that file. Make sure it is a standard, text-based PDF."

    else:
        # Reset UI to original state on refresh/open by clearing transient session state
        for key in ("events", "recurring"):
            session.pop(key, None)

    # Only group dated events for monthly view
    grouped = group_events_by_month(events)

    return render_template(
        "index.html",
        grouped_events=grouped,
        recurring_items=session.get("recurring", []),
        error=error,
    )


@app.route("/save-entry", methods=["POST"])
def save_entry():
    stored = session.get("events")
    recurring = session.get("recurring")
    if not stored and not recurring:
        flash("Nothing to save yet. Upload and extract first.", "warning")
        return redirect(url_for("index"))

    visitor_id = _get_or_create_visitor_id()
    label = (request.form.get("label") or "").strip() or None
    term_year = session.get("term_year")
    start_month = session.get("start_month")
    end_month = session.get("end_month")

    events_json = json.dumps(stored or [])
    recurring_json = json.dumps(recurring or [])

    with _db_conn() as conn:
        conn.execute(
            """
            INSERT INTO saved_entry (visitor_id, created_at, term_year, start_month, end_month, events_json, recurring_json, label)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                visitor_id,
                _utcnow_iso(),
                term_year,
                start_month,
                end_month,
                events_json,
                recurring_json,
                label,
            ),
        )

    flash("Entry saved.", "info")
    resp = redirect(url_for("history"))
    # Ensure cookie is set on redirect
    if not request.cookies.get("visitor_id"):
        resp.set_cookie("visitor_id", visitor_id, httponly=True, samesite="Lax", max_age=60*60*24*365*2)
    return resp


@app.route("/history")
def history():
    visitor_id = _get_or_create_visitor_id()
    with _db_conn() as conn:
        rows = conn.execute(
            "SELECT id, created_at, term_year, start_month, end_month, events_json, label FROM saved_entry WHERE visitor_id=? ORDER BY created_at DESC",
            (visitor_id,),
        ).fetchall()

    entries = []
    for (eid, created_at, term_year, sm, em, events_json, label) in rows:
        try:
            count = len(json.loads(events_json or "[]"))
        except Exception:
            count = 0
        entries.append({
            "id": eid,
            "created_at": created_at,
            "term_year": term_year,
            "start_month": sm,
            "end_month": em,
            "label": label or "",
            "count": count,
        })

    return render_template("history.html", entries=entries)


@app.route("/history/<int:entry_id>")
def history_detail(entry_id: int):
    visitor_id = _get_or_create_visitor_id()
    with _db_conn() as conn:
        row = conn.execute(
            "SELECT id, visitor_id, created_at, term_year, start_month, end_month, events_json, recurring_json, label FROM saved_entry WHERE id=?",
            (entry_id,),
        ).fetchone()
    if not row or row[1] != visitor_id:
        flash("Entry not found.", "error")
        return redirect(url_for("history"))

    _, _, created_at, term_year, sm, em, events_json, recurring_json, label = row
    try:
        raw_events = json.loads(events_json or "[]")
    except Exception:
        raw_events = []
    try:
        recurring_items = json.loads(recurring_json or "[]")
    except Exception:
        recurring_items = []

    # Rehydrate dated events to datetime and add calendar links
    dated_events = []
    for item in raw_events:
        try:
            dt = datetime.fromisoformat(item["date"]).replace(tzinfo=None)
            dated_events.append({"date": dt, "description": item["description"]})
        except Exception:
            continue

    events_with_links = add_calendar_links(dated_events)
    grouped = group_events_by_month(events_with_links)

    return render_template(
        "history_detail.html",
        label=label or "",
        created_at=created_at,
        grouped_events=grouped,
        recurring_items=recurring_items,
    )


@app.route("/delete-entry/<int:entry_id>", methods=["POST"])
def delete_entry(entry_id: int):
    visitor_id = _get_or_create_visitor_id()
    with _db_conn() as conn:
        row = conn.execute("SELECT visitor_id FROM saved_entry WHERE id=?", (entry_id,)).fetchone()
        if not row or row[0] != visitor_id:
            flash("Entry not found.", "error")
            return redirect(url_for("history"))
        conn.execute("DELETE FROM saved_entry WHERE id=?", (entry_id,))
    flash("Entry deleted.", "info")
    return redirect(url_for("history"))


@app.route("/download-ics", methods=["POST"])
def download_ics():
    stored = session.get("events")
    if not stored:
        flash("Upload a syllabus first to generate an ICS file.", "warning")
        return redirect(url_for("index"))

    ics = build_ics(
        [
            {"date": datetime.fromisoformat(item["date"]), "description": item["description"]}
            for item in stored
        ]
    )
    response = make_response(ics)
    response.headers["Content-Disposition"] = "attachment; filename=syllabus_due_dates.ics"
    response.headers["Content-Type"] = "text/calendar; charset=utf-8"
    return response


@app.errorhandler(413)
def _file_too_large(_e):
    # Don't render index.html here: it reads request.form, which would re-parse
    # the oversized body and re-raise the 413 inside this handler.
    max_mb = app.config["MAX_CONTENT_LENGTH"] // (1024 * 1024)
    flash(f"File too large. The maximum upload size is {max_mb} MB.", "error")
    return redirect(url_for("index"))


@app.errorhandler(429)
def _rate_limited(e):
    return render_template(
        "index.html",
        grouped_events={},
        recurring_items=[],
        error=f"Too many extraction requests ({e.description}). Please try again later.",
    ), 429


# Scanned PDFs yield little or no text from the text layer; below this
# threshold we attempt OCR instead.
_OCR_MIN_TEXT_CHARS = 100


def extract_text(uploaded_file) -> str:
    """Extract raw text from an uploaded PDF, DOCX, or text file."""
    filename = uploaded_file.filename.lower()
    if filename.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")

    if filename.endswith(".docx"):
        return _extract_docx_text(uploaded_file.stream)

    if not filename.endswith(".pdf"):
        raise RuntimeError("Please upload a PDF, DOCX, or plain text file.")

    if PdfReader is None:
        raise RuntimeError("Install the 'pypdf' package to enable PDF parsing: pip install pypdf")

    uploaded_file.stream.seek(0)
    reader = PdfReader(uploaded_file.stream)
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    text = "\n".join(pages)

    if len(text.strip()) < _OCR_MIN_TEXT_CHARS:
        ocr_text = _ocr_pdf_text(uploaded_file.stream)
        if ocr_text.strip():
            logger.info("PDF text layer was empty; used OCR fallback (%d chars)", len(ocr_text))
            return ocr_text
        raise RuntimeError(
            "This PDF has no extractable text (it may be scanned). "
            "Enable OCR support (see README) or upload a text-based export."
        )
    return text


def _extract_docx_text(stream) -> str:
    if docx is None:
        raise RuntimeError("Install the 'python-docx' package to enable DOCX parsing: pip install python-docx")
    stream.seek(0)
    document = docx.Document(stream)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _ocr_pdf_text(stream) -> str:
    """OCR a PDF via pdf2image + pytesseract. Returns "" if OCR is unavailable."""
    try:
        import pytesseract  # type: ignore
        from pdf2image import convert_from_bytes  # type: ignore
    except Exception:
        return ""
    try:
        stream.seek(0)
        images = convert_from_bytes(stream.read())
        return "\n".join(pytesseract.image_to_string(image) for image in images)
    except Exception:
        logger.exception("OCR fallback failed")
        return ""


# ----- Term-window date logic -----

_MONTH_NUMBERS = {
    "jan": 1, "january": 1,
    "feb": 2, "february": 2,
    "mar": 3, "march": 3,
    "apr": 4, "april": 4,
    "may": 5,
    "jun": 6, "june": 6,
    "jul": 7, "july": 7,
    "aug": 8, "august": 8,
    "sep": 9, "sept": 9, "september": 9,
    "oct": 10, "october": 10,
    "nov": 11, "november": 11,
    "dec": 12, "december": 12,
}

_MONTH_DAY_RE = re.compile(r"([A-Za-z]+)\.?\s+(\d{1,2})")


def window_bounds(start_month: int, end_month: int, year: int) -> Tuple[datetime, datetime]:
    """Return inclusive datetime bounds for a term window.

    A window whose end month precedes its start month (e.g. August-January)
    wraps into the following year.
    """
    start = datetime(year, start_month, 1)
    end_year = year + 1 if end_month < start_month else year
    if end_month == 12:
        end = datetime(end_year + 1, 1, 1) - timedelta(days=1)
    else:
        end = datetime(end_year, end_month + 1, 1) - timedelta(days=1)
    return start, end


def in_window(d: datetime, start: datetime, end: datetime) -> bool:
    return start <= d <= end


def resolve_date_str(text: str, start_month: int, end_month: int, year: int) -> datetime | None:
    """Resolve a month-name/day string (e.g. "Jan 6") against a term window.

    In a wrapped window (August-January), months before the start month
    resolve to the following year. Invalid dates (e.g. "Feb 30") return None.
    """
    m = _MONTH_DAY_RE.search(text)
    if not m:
        return None
    month = _MONTH_NUMBERS.get(m.group(1).lower())
    if month is None:
        return None
    day = int(m.group(2))
    resolved_year = year + 1 if end_month < start_month and month < start_month else year
    try:
        return datetime(resolved_year, month, day)
    except ValueError:
        return None


def extract_due_dates(
    text: str,
    user_year: int | None = None,
    start_month: int | None = None,
    end_month: int | None = None,
    refresh_cache: bool = False,
) -> List[Dict[str, Any]]:
    """Use the LLM to extract due dates and recurring items from the syllabus text.

    Returns a list of dicts with keys:
      - "date": datetime | None
      - "description": str
    """
    # Build a recall-first candidate text and use chunking + cache-aware extraction
    candidate = build_candidate_text(text)
    # If the candidate is too small, fallback to full text to avoid missing anything
    snippet = text if len(candidate) < 1500 else candidate
    if text:
        logger.info(
            "Candidate selection: %d -> %d chars (%.0f%% reduction)",
            len(text),
            len(snippet),
            100.0 * (1 - len(snippet) / len(text)),
        )

    raw_items = _extract_with_chatgpt(
        snippet,
        user_year=user_year,
        cache_params={
            "year": user_year,
            "start_month": start_month,
            "end_month": end_month,
        },
        refresh_cache=refresh_cache,
    )
    events: List[Dict[str, Any]] = []

    window_wraps = (
        start_month is not None and end_month is not None and end_month < start_month
    )
    bounds = None
    if user_year is not None and start_month is not None and end_month is not None:
        bounds = window_bounds(start_month, end_month, user_year)

    for item in raw_items:
        title = str(item.get("title", "")).strip() or "Due"
        # Allow due_date to be None
        due_date_val = item.get("due_date")
        recurrence_val = item.get("recurrence")
        description = str(item.get("description", title)).strip()

        # If explicit due_date provided, parse it into datetime
        if isinstance(due_date_val, str) and re.fullmatch(r"\d{4}-\d{2}-\d{2}", due_date_val.strip()):
            try:
                dt = datetime.strptime(due_date_val.strip(), "%Y-%m-%d")
            except Exception:
                # Skip invalid date strings
                continue
            # Enforce authoritative user_year if provided; in a wrapped window
            # (e.g. Aug-Jan), months before the start month belong to the next year.
            if user_year is not None:
                target_year = user_year
                if window_wraps and dt.month < start_month:
                    target_year = user_year + 1
                try:
                    dt = dt.replace(year=target_year)
                except ValueError:
                    # Invalid date in target year (e.g., Feb 29 on non-leap year)
                    continue
            # Drop dates outside the user-selected term window
            if bounds is not None and not in_window(dt, *bounds):
                continue
            events.append({"date": dt, "description": f"{title}: {description}"})
        else:
            # No explicit date. If recurrence exists, keep as undated recurring item.
            rec_str = (str(recurrence_val).strip() if recurrence_val is not None else "")
            if rec_str:
                events.append({"date": None, "description": f"{title}: {rec_str}"})
            # If neither due_date nor recurrence, ignore

    # Deduplicate and sort
    seen = set()
    unique: List[Dict[str, Any]] = []
    # Sort with None dates last, keep deterministic order
    def _sort_key(e: Dict[str, Any]):
        return (e["date"] is None, e["date"] or datetime.max)
    for e in sorted(events, key=_sort_key):
        if e["date"] is None:
            sig = (None, e["description"].lower())
        else:
            sig = (e["date"].date().isoformat(), e["description"].lower())
        if sig in seen:
            continue
        seen.add(sig)
        unique.append(e)

    logger.debug("Events kept: %d", len(unique))
    return unique


def add_calendar_links(events: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    from urllib.parse import quote

    enhanced = []
    for event in events:
        date = event["date"]
        description = event["description"].strip()
        start = date.strftime("%Y%m%d")
        end = (date + timedelta(days=1)).strftime("%Y%m%d")
        base_url = "https://calendar.google.com/calendar/render?action=TEMPLATE"
        link = f"{base_url}&text={quote(description)}&dates={start}/{end}&details={quote('Imported from syllabus calendar')}"
        enhanced.append({"date": date, "description": description, "google_link": link})
    return enhanced


def group_events_by_month(events: List[Dict[str, Any]]):
    grouped = {}
    for event in events:
        key = event["date"].strftime("%B %Y")
        grouped.setdefault(key, []).append(event)
    for event_list in grouped.values():
        event_list.sort(key=lambda e: e["date"])
    return dict(sorted(grouped.items(), key=lambda item: datetime.strptime(item[0], "%B %Y")))


def build_ics(events: List[Dict[str, Any]]) -> str:
    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Syllabus Calendar//EN",
    ]
    for event in events:
        date = event["date"].strftime("%Y%m%d")
        end_date = (event["date"] + timedelta(days=1)).strftime("%Y%m%d")
        description = _escape_ics_text(event["description"].replace("\n", " "))
        uid = f"{date}-{abs(hash(description))}@syllabus-calendar"
        lines.extend(
            [
                "BEGIN:VEVENT",
                f"UID:{uid}",
                f"DTSTAMP:{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}",
                f"DTSTART;VALUE=DATE:{date}",
                f"DTEND;VALUE=DATE:{end_date}",
                f"SUMMARY:{description}",
                "END:VEVENT",
            ]
        )
    lines.append("END:VCALENDAR")
    return "\r\n".join(lines) + "\r\n"


def _escape_ics_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace(";", "\\;").replace(",", "\\,")


DATE_WORDS = r"""
    jan|january|feb|february|mar|march|apr|april|may|jun|june|jul|july|aug|august|sep|sept|september|oct|october|nov|november|dec|december|
    mon|tue|tues|wed|thu|thur|thurs|fri|sat|sun|monday|tuesday|wednesday|thursday|friday|saturday|sunday|
    \b\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?|\b\d{4}-\d{2}-\d{2}|\b\d{1,2}:\d{2}\s?(?:am|pm)|\b\d{1,2}(?:am|pm)
""".replace("\n"," ")

KEYWORDS = [
    "due", "deadline", "submit", "submission", "turn in", "upload",
    "exam", "test", "quiz", "midterm", "final",
    "assignment", "homework", "hw", "project", "paper", "essay", "lab", "report", "proposal", "draft",
    "presentation", "checkpoint", "milestone", "gradescope", "canvas", "turnitin",
]


def build_candidate_text(full_text: str) -> str:
    """Recall-first selector: choose lines with deadline cues or date-like tokens, with ±3 lines context.

    Preserves order and deduplicates overlapping ranges. Falls back to full text if too small.
    """
    if not full_text:
        return ""

    lines = full_text.splitlines()
    matches: List[Tuple[int, int]] = []
    date_rx = re.compile(DATE_WORDS, re.I)
    kw_rx = re.compile("|".join(re.escape(k) for k in KEYWORDS), re.I)

    for i, line in enumerate(lines):
        text = line.strip()
        if not text:
            continue
        if kw_rx.search(text) or date_rx.search(text):
            start = max(0, i - 3)
            end = min(len(lines), i + 4)  # exclusive
            matches.append((start, end))

    if not matches:
        return full_text

    # Merge overlapping ranges
    matches.sort()
    merged: List[Tuple[int, int]] = []
    cur_s, cur_e = matches[0]
    for s, e in matches[1:]:
        if s <= cur_e:
            cur_e = max(cur_e, e)
        else:
            merged.append((cur_s, cur_e))
            cur_s, cur_e = s, e
    merged.append((cur_s, cur_e))

    parts = []
    for s, e in merged:
        parts.extend(lines[s:e])
        parts.append("")  # spacer to avoid run-on
    candidate = "\n".join(parts).strip()

    # If very small, better to send the full text to avoid misses
    return candidate if len(candidate) >= 1500 else full_text


# Strict JSON Schema for OpenAI structured outputs: the model is constrained to
# return {"items": [...]} matching this shape, eliminating free-form JSON parsing.
_EXTRACTION_SCHEMA = {
    "name": "syllabus_deadlines",
    "strict": True,
    "schema": {
        "type": "object",
        "properties": {
            "items": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "title": {"type": "string"},
                        "due_date": {
                            "type": ["string", "null"],
                            "description": "Explicit date from the text as YYYY-MM-DD, else null",
                        },
                        "recurrence": {"type": ["string", "null"]},
                        "description": {"type": "string"},
                    },
                    "required": ["title", "due_date", "recurrence", "description"],
                    "additionalProperties": False,
                },
            }
        },
        "required": ["items"],
        "additionalProperties": False,
    },
}


def _call_llm(client, model: str, system: str, user: str) -> str:
    messages = [{"role": "system", "content": system}, {"role": "user", "content": user}]
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0,
            response_format={"type": "json_schema", "json_schema": _EXTRACTION_SCHEMA},
        )
    except TypeError:
        # SDK too old to know response_format
        logger.warning("Structured outputs unavailable in this SDK; falling back to free-form JSON")
        resp = client.chat.completions.create(model=model, messages=messages, temperature=0)
    except Exception as exc:
        # Models without structured-output support reject the request
        if type(exc).__name__ != "BadRequestError":
            raise
        logger.warning("Model %s rejected structured outputs; falling back to free-form JSON", model)
        resp = client.chat.completions.create(model=model, messages=messages, temperature=0)
    return (resp.choices[0].message.content or "").strip()


def _parse_llm_content(content: str) -> List[Dict[str, Any]]:
    """Parse LLM output into a list of item dicts.

    Handles the structured-output object form ({"items": [...]}), a bare JSON
    array, fenced code blocks, and JSON embedded in prose. Returns [] on garbage.
    """
    content = content.strip()
    if content.startswith("```"):
        content = re.sub(r"^```(?:json)?", "", content).strip()
        if content.endswith("```"):
            content = content[:-3].strip()

    try:
        data = json.loads(content)
    except Exception:
        m = re.search(r"\[.*\]", content, flags=re.S)
        if not m:
            return []
        try:
            data = json.loads(m.group(0))
        except Exception:
            return []

    if isinstance(data, dict):
        data = data.get("items", [])
    if isinstance(data, list):
        return [d for d in data if isinstance(d, dict)]
    return []


def _extract_with_chatgpt(
    text: str,
    user_year: int | None = None,
    cache_params: Dict[str, Any] | None = None,
    refresh_cache: bool = False,
) -> List[Dict[str, Any]]:
    params = cache_params or {}
    key_basis = json.dumps({"text": text, "params": params, "year": user_year}, ensure_ascii=False, sort_keys=True)
    cache_key = hashlib.sha256(key_basis.encode("utf-8")).hexdigest()

    if not refresh_cache:
        with _db_conn() as conn:
            row = conn.execute("SELECT events_json FROM extraction_cache WHERE cache_key=?", (cache_key,)).fetchone()
        if row:
            try:
                result = json.loads(row[0])
                logger.info("Extraction cache hit (%s…)", cache_key[:12])
                return result
            except Exception:
                pass

    try:
        from openai import OpenAI  # type: ignore
    except Exception as exc:
        raise RuntimeError("The 'openai' package is not installed: pip install openai") from exc

    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set. Add it to .env or your environment.")

    client = OpenAI(api_key=api_key)
    model = os.environ.get("OPENAI_DEADLINE_MODEL", "gpt-4o-mini")

    max_chars = int(os.environ.get("SYLLABUS_MAX_CHARS", "24000"))
    chunks = [text] if len(text) <= max_chars else [text[i:i+max_chars] for i in range(0, len(text), max_chars)]

    target_year = str(user_year) if user_year is not None else "UNKNOWN"

    system = (
        "You extract graded deadlines and due dates. "
        'Output ONLY a valid JSON object of the form { "items": [ ... ] } (no prose). '
        "Never infer or guess calendar dates from month headers, semester ranges, or context. "
        "Only include a due_date if there is an explicit date string present in the text. "
        "If an item is recurring and no explicit calendar date is given, set due_date to null and include recurrence."
    )

    merged: List[Dict[str, Any]] = []
    t0 = time.perf_counter()

    for chunk in chunks:
        user = (
            "Extract graded deadlines from this syllabus.\n\n"
            f"User-selected academic year: {target_year} (authoritative; override conflicting years in text).\n\n"
            "Rules:\n"
            '- Output ONLY a valid JSON object: { "items": [ ... ] }.\n'
            f"- If user-selected year is {target_year}, normalize dates to that year when month/day are explicit.\n"
            "- Never invent missing month/day.\n"
            "- Recurring without explicit date => due_date=null and recurrence filled.\n\n"
            "Schema per item:\n"
            '{ "title": string, "due_date": "YYYY-MM-DD" | null, "recurrence": string | null, "description": string }\n\n'
            f"Syllabus text:\n```\n{chunk}\n```"
        )

        content = _call_llm(client, model, system, user)
        merged.extend(_parse_llm_content(content))

    logger.info(
        "LLM extraction: %d chunk(s), %d chars in %.2fs (model=%s)",
        len(chunks), len(text), time.perf_counter() - t0, model,
    )

    # dedupe
    seen = set()
    deduped = []
    for d in merged:
        key = (
            d.get("due_date") or None,
            (d.get("recurrence") or "").strip().lower(),
            (d.get("title") or "").strip().lower(),
            (d.get("description") or "").strip().lower(),
        )
        if key in seen:
            continue
        seen.add(key)
        deduped.append(d)

    result = deduped

    with _db_conn() as conn:
        conn.execute(
            "INSERT OR REPLACE INTO extraction_cache (cache_key, created_at, params_json, text_len, events_json) VALUES (?, ?, ?, ?, ?)",
            (
                cache_key,
                _utcnow_iso(),
                json.dumps(params, sort_keys=True),
                len(text),
                json.dumps(result),
            ),
        )

    return result


def split_events(events: List[Dict[str, Any]]):
    """Split events into dated and recurring (undated) lists.

    Returns (dated_events, recurring_items)
    """
    dated: List[Dict[str, Any]] = []
    recurring: List[Dict[str, Any]] = []
    for e in events:
        if e.get("date") is None:
            recurring.append(e)
        else:
            dated.append(e)
    return dated, recurring


if __name__ == "__main__":
    app.run(debug=True)
